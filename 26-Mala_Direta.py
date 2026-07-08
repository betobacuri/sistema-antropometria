#<
# /* ============================================================================= *\
# Nome:              Sistema de Avaliação Antropométrica
# Arquivo:           30-Mala_Direta.py
# Descrição:         Efetua o cálculo do paciente em tempo real, otimizando o tempo no 
#                    atendimento e trabalho
# Versão:            3.3
# Data:              27/06/2026
# Observações:       Script em execução no GitHub, com a conta: betobacuri
#                    Link de acesso: https://sistema-antropometria-ap.streamlit.app/
#
# Modificação:       Data: 26/06/2026
#                    Correção do cálculo de IMC para compatibilidade com altura em cm
#                    Data: 27/06/2026
#                    Alteração da fórmula do CB
#
# Desenvolvido por:  Beto Schmitt
# /* ============================================================================= *\
#>

# Library
import os
import streamlit as st
from docx import Document
from docx.shared import RGBColor


# --- 0. CONFIGURAÇÃO DA PÁGINA ---
st.set_page_config(page_title="Sistema Antropometria", page_icon="🍏", layout="wide")

st.title("🍏 Sistema de Avaliação Antropométrica")
st.subheader("Nutricionista - Ana Paula M. Cardoso Schmitt")
st.write("---")

# --- 0.1 ESTILIZAÇÃO CUSTOMIZADA DE FONTES (CSS) ---
st.markdown(
    """
    <style>
    /* =================================================================
       1. CAIXAS DE SELEÇÃO (SELECTBOX E MULTISELECT)
       ================================================================= */
    div[data-testid="stSelectbox"] label p,
    div[data-testid="stMultiSelect"] label p {
        font-size: 24px !important;
        font-weight: bold !important;
        color: #FFFFFF !important;
        line-height: 1.5 !important;
    }
    
    div[data-testid="stSelectbox"] div[data-baseweb="select"] div {
        font-size: 20px !important;
        color: #FFFFFF !important;
    }
    
    div[data-testid="stMultiSelect"] span[data-baseweb="tag"] span {
        font-size: 18px !important;
        color: #000000 !important;
    }
    
    div[data-testid="stSelectbox"] div[data-baseweb="select"],
    div[data-testid="stMultiSelect"] div[data-baseweb="select"] {
        min-height: 54px !important;
        display: flex !important;
        align-items: center !important;
    }
    
    div[data-baseweb="popover"] ul li {
        font-size: 20px !important;
        line-height: 1.6 !important;
        padding-top: 10px !important;
        padding-bottom: 10px !important;
    }
    
    /* Esconde a opção "Select all" do multiselect */
    div[data-baseweb="popover"] ul li:first-child,
    div[data-baseweb="popover"] ul[role="listbox"] li:first-child,
    li[id*="option-select-all"] {
        display: none !important;
        height: 0px !important;
        padding: 0px !important;
        margin: 0px !important;
        overflow: hidden !important;
    }

    /* =================================================================
       2. CAMPOS DE TEXTO E NÚMERO (INPUTS) - AGORA PADRONIZADOS!
       ================================================================= */
    div[data-testid="stTextInput"] label p,
    div[data-testid="stNumberInput"] label p {
        font-size: 24px !important; 
        font-weight: bold !important; 
        color: #FFFFFF !important;
        line-height: 1.5 !important;
    }

    div[data-testid="stTextInput"] input,
    div[data-testid="stNumberInput"] input {
        font-size: 20px !important;
        color: #FFFFFF !important;
        height: 54px !important;
    }

    div[data-testid="stNumberInput"] button {
        height: 52px !important;
    }
    
    /* =================================================================
       3. AJUSTE DE BOTÕES
       ================================================================= */
    div[data-testid="stButton"] button p,
    div[data-testid="stDownloadButton"] button p {
        font-size: 20px !important;
        font-weight: bold !important;
    }
    
    div[data-testid="stButton"] button,
    div[data-testid="stDownloadButton"] button {
        padding: 10px 24px !important;
        height: auto !important;
        width: 100% !important; 
    }
    </style>
    """,
    unsafe_allow_html=True
)

# --- INICIALIZAÇÃO CONTROLADA DO SESSION STATE ---
if "key_reset" not in st.session_state:
    st.session_state.key_reset = 0
if "diagnostico_calculado" not in st.session_state:
    st.session_state.diagnostico_calculado = False
if "dados_laudo" not in st.session_state:
    st.session_state.dados_laudo = {}

# Guarda o sufixo atual para as chaves dos inputs
kr = st.session_state.key_reset

# --- 1. FUNÇÕES DE CÁLCULO ---
# Função Cálculo - Amputado
def calcular_peso_ideal_amputado(peso_ideal_inicial, lista_membros):
    if not lista_membros:
        return peso_ideal_inicial, 0.0
    tabela_amputacao = {
        "Antebraço Direito": 2.3, "Antebraço Esquerdo": 2.3, "Braço Direito até o ombro": 6.5, 
        "Braço Esquerdo até o ombro": 6.5, "Mão Direita": 0.8, "Mão Esquerda": 0.8, 
        "Perna Direita abaixo do joelho": 5.3, "Perna Direita acima do joelho": 11.6, 
        "Perna Direita inteira": 16.0, "Perna Esquerda abaixo do joelho": 5.3,
        "Perna Esquerda acima do joelho": 11.6, "Perna Esquerda inteira": 16.0,
        "Pé Direito": 1.5, "Pé Esquerda": 1.5
    }
    total_percentual = sum(tabela_amputacao.get(membro, 0.0) for membro in lista_membros)
    if total_percentual > 100.0: total_percentual = 100.0
    return peso_ideal_inicial * (1 - (total_percentual / 100)), total_percentual

# Função Cálculo - Edema
def obter_desconto_edema(grau_edema):
    return {"Sem Edema": 0.0, "Grau I": 1.0, "Grau II": 2.5, "Grau III": 5.0, "Grau IV": 10.0}.get(grau_edema, 0.0)

# Função Cálculo - Ascite
def obter_desconto_ascite(grau_ascite):
    return {"Sem Ascite": 0.0, "Leve": 2.2, "Moderada": 6.0, "Grave": 14.0}.get(grau_ascite, 0.0)

# --- 2. CAPTURA DE DADOS (ENTRADAS COM CHAVES DINÂMICAS) ---
st.markdown("### 📋 Dados do Paciente e Antropometria")

# Tela Descrição dos Pacientes 
col1, col2, col3, col4 = st.columns(4)

# Coluna 1
with col1:
    nome_paciente = st.text_input("Nome do Paciente", value="", key=f"nome_{kr}", placeholder="Digite o nome...")
    
    # Idade como texto para iniciar em branco
    f_idade_txt = st.text_input("Idade", value="", key=f"idade_{kr}", placeholder="Ex: 35")
    f_idade = int(f_idade_txt) if f_idade_txt.strip().isdigit() else 0
    f_genero = st.selectbox("Sexo", ["Selecione...", "Masculino", "Feminino"], key=f"genero_{kr}")
    f_etnia = st.selectbox("Raça/Etnia", ["Selecione...", "Negra", "Branca"], key=f"etnia_{kr}")

# Função auxiliar para converter o texto em float aceitando ponto ou vírgula
def converter_para_float(valor_txt):
    if not valor_txt.strip():
        return 0.0
    try:
        return float(valor_txt.replace(',', '.'))
    except ValueError:
        return 0.0

# Coluna 2
with col2:
    f_peso_atual_txt = st.text_input("Peso Atual Aferido", value="", key=f"peso_atual_{kr}", placeholder="Ex: 75.5")
    f_peso_atual = converter_para_float(f_peso_atual_txt)
    f_peso_habitual_txt = st.text_input("Peso Habitual", value="", key=f"peso_hab_{kr}", placeholder="Ex: 80.0")
    f_peso_habitual = converter_para_float(f_peso_habitual_txt)
    f_altura_txt = st.text_input("Altura (cm)", value="", key=f"altura_{kr}", placeholder="Ex: 170")
    f_altura = int(f_altura_txt) if f_altura_txt.strip().isdigit() else 0

# Coluna 3
with col3:
    f_cb_txt = st.text_input("Circunferência do Braço (cm)", value="", key=f"cb_{kr}", placeholder="Ex: 32.5")
    f_cb = converter_para_float(f_cb_txt)
    f_cp_txt = st.text_input("Circunferência da Panturrilha (cm)", value="", key=f"cp_{kr}", placeholder="Ex: 34.0")
    f_cp = converter_para_float(f_cp_txt)
    f_aj_txt = st.text_input("Altura do Joelho (cm)", value="", key=f"aj_{kr}", placeholder="Ex: 52.0")
    f_aj = converter_para_float(f_aj_txt)

# Coluna 4
with col4:
    f_status_edema = st.selectbox("Status Edema", ["Sem Edema", "Grau I", "Grau II", "Grau III", "Grau IV"], key=f"edema_{kr}")
    f_status_ascite = st.selectbox("Status Ascite", ["Sem Ascite", "Leve", "Moderada", "Grave"], key=f"ascite_{kr}")
    
    opcoes_amputacao = [        
        "Antebraço Direito", "Antebraço Esquerdo", "Braço Direito até o ombro", 
        "Braço Esquerdo até o ombro", "Mão Direita", "Mão Esquerda", 
        "Perna Direita abaixo do joelho", "Perna Direita acima do joelho", 
        "Perna Direita inteira", "Perna Esquerda abaixo do joelho",
        "Perna Esquerda acima do joelho", "Perna Esquerda inteira", 
        "Pé Direito", "Pé Esquerdo"
    ]
    f_status_amputado = st.multiselect("Corpo Integro (Clique Aqui Para Alterar)", opcoes_amputacao, key=f"amputado_{kr}", placeholder="Selecione...")

f_peso_ideal = f_peso_atual  
st.write("---")

# --- 3. BOTÃO: CALCULAR DIAGNÓSTICO ---
if st.button("📊 Calcular Diagnóstico", type="secondary", key="btn_calcular"):
    if not nome_paciente.strip() or f_altura == 0 or f_peso_atual == 0.0:
        st.error("⚠️ Por favor, preencha o Nome, Peso e Altura do paciente antes de calcular!")
    else:
        dados = {}
        peso_pos_amputacao, pct_membro = calcular_peso_ideal_amputado(f_peso_atual, f_status_amputado)
        kg_desconto_edema = obter_desconto_edema(f_status_edema)
        kg_desconto_ascite = obter_desconto_ascite(f_status_ascite)
        
        peso_corrigido = peso_pos_amputacao - kg_desconto_edema - kg_desconto_ascite
        if peso_corrigido < 0: peso_corrigido = 0.0  
        
        # CORREÇÃO: Altura convertida de cm para metros para cálculo correto do IMC
        imc_real = peso_corrigido / ((f_altura / 100) ** 2) if f_altura > 0 else 0.0
        
        if f_status_amputado:  
            membros_texto = ", ".join(f_status_amputado)
            st.warning(f"⚠️ Amputação ({membros_texto}): -{pct_membro:.1f}%")
            
        if f_status_edema != "Sem Edema" or f_status_ascite != "Sem Ascite":
            st.info(f"💧 Retenção Hídrica: Descontado {kg_desconto_edema}kg (Edema) e {kg_desconto_ascite}kg (Ascite).")

        dados['Nome'] = nome_paciente
        dados['F_Idade'] = str(f_idade)
        dados['F_Etnia'] = f_etnia
        dados['F_CB-(cm)'] = f"{f_cb:.1f}"
        dados['F_CP-(cm)'] = f"{f_cp:.1f}"
        dados['F_AJ-(cm)'] = f"{f_aj:.1f}"
        dados['F_Peso_Atual-(aferido)'] = f"{peso_corrigido:.1f}"
        dados['F_Peso_Habitual-(referido)'] = f"{f_peso_habitual:.1f}"
        # Convertendo para metros apenas na string de exibição se o laudo esperar em metros, ou mantendo original:
        dados['F_Altura-(referida)'] = f"{(f_altura/100):.2f}"
        dados['F_IMC_Dados_Refer'] = f"{imc_real:.2f}"
        dados['F_IMC_Est_Altura'] = f"{(f_altura/100):.2f}"
        dados['F_IMC_Est_Peso'] = f"{peso_corrigido:.1f}"
        dados['F_IMC_Est_IMC'] = f"{imc_real:.2f}"
        
        if f_idade >= 60:
            dados['F_IMC_Dados_Refer_Adulto'] = "—"
            dados['F_IMC_Est_Adulto'] = ""
            if imc_real < 21.99: classe_imc = "Baixo peso"
            elif imc_real <= 23.99: classe_imc = "Risco de déficit"
            elif imc_real <= 26.99: classe_imc = "Eutrofia"
            else: classe_imc = "Sobrepeso"
            dados['F_IMC_Dados_Refer_Idoso'] = classe_imc
            dados['F_IMC_Est_Idoso'] = classe_imc
        else:
            dados['F_IMC_Dados_Refer_Idoso'] = "—"
            dados['F_IMC_Est_Idoso'] = ""
            if imc_real < 18.49: classe_imc = "Magreza grau I"
            elif imc_real <= 24.99: classe_imc = "Eutrofia"
            else: classe_imc = "Sobrepeso"
            dados['F_IMC_Dados_Refer_Adulto'] = classe_imc
            dados['F_IMC_Est_Adulto'] = classe_imc

        if f_cp < 31.0: dados['F_Circ_Panturrilha'] = "Baixa massa muscular"
        else: dados['F_Circ_Panturrilha'] = "Massa muscular adequada"

        # --- CÁLCULO DINÂMICO DO P50 DA CB (IGUAL AO EXCEL DO HOSPITAL) ---
        p50_cb = 0.0
        
        if f_genero == "Masculino":
            if f_idade < 18.9: p50_cb = 29.7
            elif f_idade < 24.9: p50_cb = 30.8
            elif f_idade < 34.9: p50_cb = 31.9
            elif f_idade < 44.9: p50_cb = 32.6
            elif f_idade < 54.9: p50_cb = 32.2
            elif f_idade < 64.9: p50_cb = 31.7
            elif f_idade < 74.9: p50_cb = 30.7
            else: p50_cb = None  # Equivale ao "Não se aplica" para >= 75 anos
            
        elif f_genero == "Feminino":
            # Caso ela precise para mulheres no futuro, mantemos o padrão atual
            p50_cb = 28.5 
        else:
            p50_cb = None

        # --- APLICAÇÃO DOS RESULTADOS NO LAUDO ---
        if p50_cb is None or p50_cb == 0.0:
            dados['F_Circ_Braco-P50'] = "Não se aplica"
            dados['F_Circ_Braco-%CB'] = "Não se aplica"
            dados['F_Circ_Braco-Classificacao'] = "Não se aplica"
        else:
            dados['F_Circ_Braco-P50'] = f"{p50_cb:.1f}"
            if f_cb > 0:
                # Faz exatamente a conta do Excel: (CB * 100) / P50
                adequacao_cb = (f_cb * 100) / p50_cb
                dados['F_Circ_Braco-%CB'] = f"{adequacao_cb:.1f}%"
                
                # Mantém a classificação automática baseada na adequação
                if adequacao_cb < 70.0: dados['F_Circ_Braco-Classificacao'] = "Desnutrição Grave"
                elif adequacao_cb <= 90.0: dados['F_Circ_Braco-Classificacao'] = "Desnutrição Leve"
                else: dados['F_Circ_Braco-Classificacao'] = "Eutrofia"
            else:
                dados['F_Circ_Braco-%CB'] = "—"
                dados['F_Circ_Braco-Classificacao'] = "—"
        dados['_peso_corrigido_interno'] = peso_corrigido
        st.session_state.dados_laudo = dados
        st.session_state.diagnostico_calculado = True

# --- 4. APRESENTAÇÃO DO DIAGNÓSTICO E REGRAS DE BOLSO ---
if st.session_state.diagnostico_calculado:
    d = st.session_state.dados_laudo

    st.markdown("### 📊 Resumo do Diagnóstico Encontrado")
    c_res1, c_res2, c_res3 = st.columns(3)
    with c_res1: st.metric("IMC Calculado", d['F_IMC_Dados_Refer'])
    with c_res2: st.metric("Classificação CB", d['F_Circ_Braco-Classificacao'])
    with c_res3: st.metric("Status Panturrilha", d['F_Circ_Panturrilha'])
        
    st.write("---")
    st.markdown("### 🍽️ Necessidades Nutricionais (Regra de Bolso)")
    col_vet, col_ptn, col_hid = st.columns(3)

    # Listagem do fator VET 
    with col_vet:
        st.markdown("**⚡ Energia (VET)**")
        opcao_vet = st.selectbox("Fator de VET (Diagnóstico)", ["Selecione...", "Cirrose Hepática", "Cirurgia Eletiva em Geral", "Eutrófico", "Ganho de Peso sem Estresse", "Manutenção de Peso sem Estresse", "Obeso (CRÍTICO) - Peso Atual", "Obeso (CRÍTICO) - Peso Ideal", "Obeso","Pacientes CRÍTICOS - Trauma", "Perda de Peso", "Politrauma", "Quimio e Radioterapia - Manutenção de Peso", "Quimio e Radioterapia - Ganho de Peso", "Quimio e Radioterapia - Obeso"], index=0, key=f"vet_sel_{kr}")
    
    # Listagem Proteína 
    with col_ptn:
        st.markdown("**🥩 Proteínas**")        
        opcao_ptn = st.selectbox("Fator de Proteína (Diagnóstico)", ["Selecione...", "Catabolismo Moderado (Cirrose Hepática, Doença de Crohn etc)", "Desnutrido, Hipercatabolismo (Pancreatite Aguda Grave, SARA etc)", "Diálise", "Queimados e com Fístula", "ELA ou EscleroseMúltipla", "Eutrófico", "Neurocríticos", "Obeso (30 a 40 kg/m²) - Peso Ideal", "Obeso (>40 kg/m²) - Peso Ideal", "Obeso Crítico (30 a 40 kg/m²) - Peso Ideal", "Obeso Crítico (>40 kg/m²) - PesoIdeal", "Paciente Oncológico - Estresse Grave", "Paciente Oncológico - Estresse Moderado", "Quimio e Radioterapia - Estresse Grave e Repleção Proteica", "Quimio e Radioterapia - Estresse Moderado", "Quimio e Radioterapia - sem Complicações"], index=0, key=f"ptn_sel_{kr}")

    # Listagem Hídrica 
    with col_hid:
        st.markdown("**💧 Recomendações Hídricas**")
        opcao_hid = st.selectbox("Fator de Líquidos (Estado)", ["Selecione...", "Estado de Hidratação Normal (Função Renal e Cardíaca)", "Pacientes Renais"], index=0, key=f"hid_sel_{kr}")

    st.write("---")

# Botão Gerar Relatório Word
    if st.button("💾 Gerar Relatório Word Final", type="primary", key="btn_gerar_word"):
        if opcao_vet == "Selecione..." or opcao_ptn == "Selecione..." or opcao_hid == "Selecione...":
            st.error("⚠️ **Atenção!** Por favor, selecione os fatores de VET, Proteína e Líquidos antes de gerar o relatório.")
        else:
            peso_base = d.get('_peso_corrigido_interno', f_peso_atual)
            
            # 1. Ajuste e padronização da tabela de VET
            tabela_vet = {
                "eutrófico": (20, 30), "perda de peso": (20, 25), "manutenção de peso sem estresse": (25, 30), 
                "ganho de peso sem estresse": (30, 35), "pacientes críticos, trauma": (20, 25), "politrauma": (40, 40), 
                "obeso": (12, 20), "obeso (crítico) - peso atual": (11, 14), "cirurgia eletiva em geral": (32, 32), 
                "cirrose hepática": (30, 35), "quimio e radioterapia - obesos": (20, 25), 
                "quimio e radioterapia - manutenção de peso": (25, 30), "quimio e radioterapia - ganho de peso": (30, 35), 
                "obeso (crítico) - peso ideal": (25, 25), "pacientes críticos - trauma": (20, 25)
            }

            # Buscando convertendo a opção do selectbox para minúsculo
            fator_min, f_max = tabela_vet.get(opcao_vet.lower(), (0, 0))
            vet_minimo_calculado = peso_base * fator_min
            vet_maximo_calculado = peso_base * f_max

            d['F_VET_Paciente_Result01'] = opcao_vet
            d['F_VET_Min_Result01'] = f"{fator_min:.1f}" if fator_min > 0 else "0"
            d['F_VET_Max_Result01'] = f"{f_max:.1f}" if f_max > 0 else "0"
            d['F_VET_Kcal-Dia01'] = f"{vet_minimo_calculado:.0f} a {vet_maximo_calculado:.0f} kcal/dia" if vet_minimo_calculado != vet_maximo_calculado else f"{vet_minimo_calculado:.0f} kcal/dia"

            # 2. Ajuste e padronização da tabela de Proteínas
            tabela_proteina = {
                "eutrófico": (1.0, 1.2), "ela ou esclerose múltipla": (0.8, 1.2), 
                "catabolismo moderado (cirrose hepática, doença de crohn etc)": (1.2, 1.5), 
                "catabolismo moderado (cirrose hepática, doença de crohn, etc)": (1.2, 1.5),
                "paciente oncológico - estresse moderado": (1.2, 1.5), "paciente oncológico - estresse grave": (1.5, 2.0), 
                "quimio e radioterapia - sem complicações": (1.0, 1.2), "quimio e radioterapia - estresse moderado": (1.2, 1.5), 
                "quimio e radioterapia - estresse grave e repleção proteica": (1.5, 2.0), "neurocríticos": (1.3, 1.5), 
                "desnutrido, hipercatabolismo (pancreatite aguda grave, sara etc)": (1.5, 2.0), 
                "desnutrido, hipercatabolismo (pancreatite aguda grave, sara, etc)": (1.5, 2.0), 
                "diálise": (1.0, 1.5), "queimados e com fístula": (2.0, 2.5), 
                "obeso (30 a 40 kg/m²) - peso ideal": (2.5, 2.5), "obeso (>40 kg/m²) - peso ideal": (3.0, 3.0), 
                "obeso crítico (30 a 40 kg/m²) - peso ideal": (2.0, 2.0), "obeso crítico (>40 kg/m²) - peso ideal": (2.5, 2.5)
            }

            ptn_min, ptn_max = tabela_proteina.get(opcao_ptn.lower(), (0.0, 0.0))
            
            # Verificação do peso ideal 
            base_calculo_ptn = f_peso_ideal if "peso ideal" in opcao_ptn.lower() else peso_base
            ptn_min_calculada = base_calculo_ptn * ptn_min
            ptn_max_calculada = base_calculo_ptn * ptn_max

            d['F_Proteina_Paciente_Result01'] = opcao_ptn
            d['F_Proteina_Paciente_Min01'] = f"{ptn_min:.1f}"
            d['F_Proteina_Paciente_Max01'] = f"{ptn_max:.1f}"
            d['F_Proteina_Recomenda_g-kg-dia01'] = f"{ptn_min_calculada:.0f} a {ptn_max_calculada:.0f} g/dia" if ptn_min_calculada != ptn_max_calculada else f"{ptn_min_calculada:.0f} g/dia"

            # 3. Ajuste e padronização da tabela Hídrica
            d['F_Hidrico_Paciente'] = opcao_hid
            if opcao_hid.lower() == "pacientes renais":
                d['F_Hidrico_ml-kg-dia01'] = "média de urina 24h + 500ml"
                d['F_Hidrico_ml-kcal01'], d['F_Hidrico_ml-kcal02'], d['F_Hidrico_ml-kg-dia03'] = "", "", ""
                d['F_Hidrico_ml-kg-dia02'] = "média de urina 24h + 500ml"
            else:
                d['F_Hidrico_ml-kg-dia01'], d['F_Hidrico_ml-kcal01'] = "30 a 40", "1,5"
                d['F_Hidrico_ml-kg-dia02'] = f"{peso_base * 30:.0f} ml"
                d['F_Hidrico_ml-kg-dia03'] = f"{peso_base * 40:.0f} ml"
                d['F_Hidrico_ml-kcal02'] = f"{vet_minimo_calculado * 1.5:.0f} ml"

            # Campos Restantes - SEM USO
            campos_restantes = ['F_Estimativa_Altura_Idoso', 'F_Estimativa_Altura_Branco', 'F_Estimativa_Altura_Negro', 'F_Estimativa_Peso_Negro', 'F_Estimativa_Peso_Negro_Idoso', 'F_Estimativa_Peso_Branco', 'F_Estimativa_Peso_Branco_Idoso', 'F_Estimativa_Peso_Negro_Edema', 'F_Estimativa_Peso_Negro_Idoso_Edema', 'F_Estimativa_Peso_Branco_Edema', 'F_Estimativa_Peso_Branco_Idoso_Edema', 'F_Estimativa_Peso_Negro_Ascite', 'F_Estimativa_Peso_Negro_Idoso_Ascite', 'F_Estimativa_Peso_Branco_Ascite', 'F_Estimativa_Peso_Branco_Idoso_Ascite', 'F_Estimativa_Peso_Negro_Amputado', 'F_Estimativa_Peso_Negro_Idoso_Amputado', 'F_Estimativa_Peso_Branco_Amputado', 'F_Estimativa_Peso_Branco_Idoso_Amputado', 'F_Perda_Peso_%PPR', 'F_Perda_Peso_%PPR_Tempo', 'F_Perda_Peso_%PPR_Result01', 'F_Perda_Peso_%PPR_Result02', 'F_Perda_Peso_%PPR_Result03', 'F_Perda_Peso_%PPR_Result04', 'F_Peso_Ideal_Idade_Adultos', 'F_Peso_Ideal_Adultos_Result', 'F_Peso_Ideal_Idade_Idoso', 'F_Peso_Ideal_Idoso_Resul01', 'F_Peso_Ideal_Idoso_Resul02']
            for campo in campos_restantes: d[campo] = "[Pendente]"

            pasta_destino = "Relatório"
            if not os.path.exists(pasta_destino): os.makedirs(pasta_destino)
                
            nome_paciente_limpo = nome_paciente.replace(' ', '_') if nome_paciente else "Paciente"
            nome_saida = os.path.join(pasta_destino, f"Laudo_{nome_paciente_limpo}.docx")
            
            if os.path.exists("modelo.docx"):
                doc = Document("modelo.docx")
                def substituir_e_formatar(paragrafo, dados_dicionario):
                    texto_modificado = paragrafo.text
                    alterou = False
                    for chave, valor in dados_dicionario.items():
                        tag = f"{{{{{chave}}}}}"
                        if tag in texto_modificado:
                            texto_modificado = texto_modificado.replace(tag, valor)
                            alterou = True
                    if alterou:
                        paragrafo.text = texto_modificado
                        for run in paragrafo.runs:
                            run.font.name = 'Arial'
                            run.font.color.rgb = RGBColor(0, 0, 0)

                for p in doc.paragraphs: substituir_e_formatar(p, d)
                for tabela in doc.tables:
                    for linha in tabela.rows:
                        for celula in linha.cells:
                            for p in celula.paragraphs: substituir_e_formatar(p, d)
                                        
                doc.save(nome_saida)
                st.success("🎉 Diagnóstico e Necessidades Nutricionais calculados!")
                st.session_state['relatorio_pronto'] = True
                st.session_state['nome_saida_salvo'] = nome_saida
                st.session_state['nome_paciente_limpo'] = nome_paciente_limpo
            else:
                st.error("❌ Erro: O arquivo 'modelo.docx' não foi encontrado.")

# --- 5. BLOCO DE DOWNLOAD E LIMPAR CAMPOS ---
if st.session_state.get('relatorio_pronto', False):
    st.markdown("---")
    with open(st.session_state['nome_saida_salvo'], "rb") as file:
        # Botão Download
        st.download_button(
            label="📥 Baixar o Laudo",
            data=file,
            file_name=f"Laudo_{st.session_state['nome_paciente_limpo']}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="btn_download_final"
        )
    
    # Botão Limpar os Campos
    if st.button("🔄 Novo Paciente", key="btn_novo_paciente"):
        # Descobre o próximo valor do ciclo antes de limpar tudo
        proximo_ciclo = st.session_state.key_reset + 1
        
        # Limpa o estado
        st.session_state.clear()
        
        # Garante que o novo ciclo sobreviva à limpeza
        st.session_state.key_reset = proximo_ciclo
        
        # Recarrega forçando IDs completamente inéditos nos inputs
        st.rerun()
  
# Fim.
